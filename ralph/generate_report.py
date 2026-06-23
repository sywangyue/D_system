#!/usr/bin/env python3
"""AM MEDICAL 背调报告生成器 — 读取 findings/ 生成中文 .docx"""
import json, os
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))

def load_json(path):
    with open(os.path.join(BASE, path)) as f:
        return json.load(f)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return h

def add_table_from_dict(doc, data, col_widths=None):
    """data: list of (key, value) tuples"""
    table = doc.add_table(rows=len(data)+1, cols=2, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = '字段'
    hdr[1].text = '内容'
    for i, (k, v) in enumerate(data):
        table.rows[i+1].cells[0].text = str(k)
        table.rows[i+1].cells[1].text = str(v)
    return table

def add_evidence_badge(paragraph, strength):
    """Add evidence badge"""
    colors = {'HIGH': 'FF4444', 'MED': 'FFAA00', 'LOW': '888888'}
    labels = {'HIGH': '■ 强证据', 'MED': '■ 中等证据', 'LOW': '■ 弱证据'}
    run = paragraph.add_run(f" {labels.get(strength, strength)} ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*tuple(int(colors.get(strength, '888888')[i:i+2], 16) for i in (0, 2, 4)))

def main():
    today = date.today().strftime('%Y%m%d')
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    
    # ---- COVER ----
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AM MEDICAL EQUIPMENT PTE. LTD.\n黑代理深度背调报告')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'\n排查日期: {date.today().strftime("%Y-%m-%d")}\n')
    meta.add_run('执行模式: Ralph 自治背调循环 (S01-S08 全量)\n')
    meta.add_run('输出格式: Word (.docx) 中文\n')
    meta.add_run('展会: MEDICA 2025 Düsseldorf + 历史届次\n')
    meta.add_run('执行方: Claude Code (Ralph Loop)\n')
    
    doc.add_page_break()
    
    # ---- PART 1: 企业一般情况 ----
    add_heading_styled(doc, '第一部分 企业一般情况', 1)
    
    # 1.1 AM MEDICAL
    add_heading_styled(doc, '1.1 AM MEDICAL EQUIPMENT PTE. LTD.（新加坡壳公司）', 2)
    
    s02 = load_json('findings/02_sg_registry.json')['findings']['current_entity']
    
    ame_data = [
        ('公司全名', s02['name']),
        ('UEN (注册号)', s02['uen']),
        ('成立日期', s02['incorporation_date']),
        ('公司类型', s02['company_type']),
        ('经营状态', s02['status']),
        ('注册资本', s02['paid_up_capital']),
        ('官员人数', str(s02['officers_count'])),
        ('注册地址', s02['registered_address']),
        ('地址类型', s02['address_type']),
        ('主营业务 (SSIC)', s02['primary_ssic']),
        ('次要业务 (SSIC)', s02['secondary_ssic']),
        ('年报提交日', s02['annual_return_date']),
        ('财报截止日', s02['account_due_date']),
    ]
    add_table_from_dict(doc, ame_data)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('⚠ 壳公司识别特征（已确认）:').bold = True
    indicators = [
        '注册资本 SGD 1,000（新加坡法定最低值）',
        '注册地址为虚拟办公室（同单元 826+ 家公司注册）',
        '次要经营范围含"公关/营销/品牌咨询"（医疗器械批发商不应有此项）',
        '企业联系邮箱为个人 Gmail (shoneyuwell@gmail.com)，非企业域名邮箱',
        '官网 sg-amgroup.com 使用中国建站平台 FaiscoBuild 搭建，底部含中文字样',
        '官网域名注册人为中国公司 奥瑞文化传媒镇江有限公司（非 AM MEDICAL 自身）',
        '前身实体 AM MEDICAL PTE. LTD. (202013171K) 已被 ACRA 除名',
        '官网宣称"成立于2010年"但新加坡注册记录为2023年（历史伪造）',
        '域名创建于2019年（早于公司注册4年），品牌策划早于法律实体',
    ]
    for ind in indicators:
        doc.add_paragraph(f'  • {ind}', style='List Bullet')
    
    # Predecessor
    add_heading_styled(doc, '1.2 前身实体: AM MEDICAL PTE. LTD.（已注销）', 2)
    pred = load_json('findings/02_sg_registry.json')['findings']['predecessor_entity']
    pred_data = [
        ('公司全名', pred['name']),
        ('UEN', pred['uen']),
        ('成立日期', pred['incorporation_date']),
        ('经营状态', pred['status']),
        ('注册资本', pred['paid_up_capital']),
        ('注册地址', pred['registered_address']),
    ]
    add_table_from_dict(doc, pred_data)
    doc.add_paragraph(f"注销原因推断: {pred['strike_off_reason']}")
    
    # 1.3 杭州仁然
    add_heading_styled(doc, '1.3 杭州仁然文化传媒有限公司（中国签约主体）', 2)
    s03 = load_json('findings/03_cn_registry.json')['findings']['primary_target']
    hrr_data = [
        ('公司全名', s03['name']),
        ('统一社会信用代码', s03['credit_code']),
        ('法定代表人', s03['legal_rep']),
        ('成立日期', s03['established']),
        ('注册资本', s03['registered_capital']),
        ('参保人数', str(s03['insured_count'])),
        ('企业类型', s03['company_type']),
        ('注册地址', s03['address']),
        ('联系电话', ', '.join(s03['contact_info']['phones'])),
        ('联系邮箱', ', '.join(s03['contact_info']['emails'])),
    ]
    add_table_from_dict(doc, hrr_data)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('⚠ 异常信号:').bold = True
    for sig in s03['anomaly_signals']:
        doc.add_paragraph(f'  • {sig}', style='List Bullet')
    
    # 赵立业
    add_heading_styled(doc, '1.4 赵立业（实际经营代表）名下其他企业', 2)
    zly_ents = load_json('findings/03_cn_registry.json')['findings']['legal_rep_other_entities']
    for ent in zly_ents:
        doc.add_paragraph(f"• {ent['name']} — {ent.get('industry', '')} — {ent.get('relevance', '')}")
    
    doc.add_page_break()
    
    # ---- PART 2: 关系图谱 ----
    add_heading_styled(doc, '第二部分 关系图谱', 1)
    
    doc.add_paragraph('以下关系图谱基于 S01-S06 全部证据链汇总。每条关联标注证据强度（HIGH/MED/LOW）。')
    doc.add_paragraph('三层结构：海外壳公司层 → 制造执行层 → 销售代理层 → 下游中国企业')
    
    # Evidence chains
    s04 = load_json('findings/04_fingerprint_links.json')['findings']
    
    add_heading_styled(doc, '2.1 已确认关联链（HIGH 强度）', 2)
    
    high_chains = [c for c in s04['chains'] if c['strength'] == 'HIGH']
    for i, chain in enumerate(high_chains, 1):
        p = doc.add_paragraph()
        p.add_run(f'证据链 {i}: {chain["title"]}').bold = True
        add_evidence_badge(p, 'HIGH')
        doc.add_paragraph(chain['description'])
        for ev in chain['evidence']:
            doc.add_paragraph(f'  • {ev["key"]}: {ev["value"]}')
        doc.add_paragraph(f'  来源: {chain["source"]}')
        doc.add_paragraph()
    
    add_heading_styled(doc, '2.2 中等证据关联链（MED 强度）', 2)
    med_chains = [c for c in s04['chains'] if c['strength'] == 'MED']
    for i, chain in enumerate(med_chains, 1):
        p = doc.add_paragraph()
        p.add_run(f'证据链 {i}: {chain["title"]}').bold = True
        add_evidence_badge(p, 'MED')
        doc.add_paragraph(chain['description'])
        doc.add_paragraph()
    
    add_heading_styled(doc, '2.3 关联网络文字解读', 2)
    
    network = s04['network_summary']
    doc.add_paragraph(f'三层结构确认:')
    doc.add_paragraph(f'  ① 海外壳公司层: {network["layer_1_shell"]}')
    doc.add_paragraph(f'  ② 制造执行层: {network["layer_2_manufacturing"]}')
    doc.add_paragraph(f'  ③ 销售代理层: {network["layer_3_sales_agent"]}')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('已证实关联:').bold = True
    for link in network['confirmed_links']:
        doc.add_paragraph(f'  ✓ {link}', style='List Bullet')
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('待进一步证实（推断）:').bold = True
    for link in network['unconfirmed_links']:
        doc.add_paragraph(f'  ○ {link}', style='List Bullet')
    
    doc.add_page_break()
    
    # ---- PART 3: 黑代理侵入流程 ----
    add_heading_styled(doc, '第三部分 黑代理侵入流程', 1)
    
    doc.add_paragraph('以下用编号步骤复现 AM MEDICAL / 杭州仁然 黑代理从注册壳公司到拆售展位的完整操作流程，每步标注证据支撑。')
    
    steps = [
        ('Step 1: 域名与品牌预置 (2019)',
         '在注册任何公司之前，先在阿里云注册域名 sg-amgroup.com，注册人为 奥瑞文化传媒镇江有限公司（任丹芳）。'
         '该域名后续成为 AM MEDICAL 的官网。品牌策划早于法律实体 4 年，说明这是一个有预谋的长期项目。',
         '证据: WHOIS 查询 (chinaz.com)，域名创建日 2019-12-04'),
        
        ('Step 2: 第一代壳公司注册 (2020-05)',
         '注册 AM MEDICAL PTE. LTD. (202013171K)，主营健康补充剂批发 + 医疗器械批发，地址为 Centro Bianco #06-01C。',
         '证据: ACRA 记录 (companies.sg)，注册日 2020-05-09'),
        
        ('Step 3: 第一代壳被除名 (~2022)',
         'AM MEDICAL PTE. LTD. 因未提交年报/财报被 ACRA 除名 (Struck Off)。壳公司生命周期结束。',
         '证据: ACRA 状态查询，状态 Struck Off'),
        
        ('Step 4: 第二代壳公司注册 (2023-06-22)',
         '注册 AM MEDICAL EQUIPMENT PTE. LTD. (202324580R)。注册资本降至 SGD 1,000（最低值），次要业务增加"公关/营销/品牌咨询"。'
         '地址为同一栋楼不同单元 (#07-02J)，推测同一企业服务商代理。',
         '证据: ACRA 记录 (companies.sg/sgbusinessdirectory.com)'),
        
        ('Step 5: 中国母集团同步注册 (2023-07-11)',
         '华尔科技（河南）集团有限公司在郑州注册成立，法人王超，经营范围含"会议及展览服务"。'
         '与 AM MEDICAL EQUIPMENT 仅隔 19 天注册——高度疑似协同设立。'
         '华尔集团旗下已有运营中的制造子公司 郑州奥利弗电子科技（2014年成立，76人参保，实际生产超声设备）。',
         '证据: QCC 报告，成立日 2023-07-11。AM MEDICAL 成立日 2023-06-22，差值 19 天。'),
        
        ('Step 6: Medica 2022-2024 期间使用旧名参展',
         '2022-2024 年以 "Am Wheelchair + Equipment Pte Ltd" (vis_hash Xoq0b3yKRgmYCYWNmjOAZA) 名义在 Medica Hall 6 D40 参展。'
         '2024 年已确认至少 2 家中国企业在同一展位（雷虎河南 + 江苏康宝），均为 high risk。',
         '证据: medica_investigation.db (2025数据) + 旧版SOP (2022-2024数据)'),
        
        ('Step 7: 中国销售壳注册 (2024-10-12)',
         '杭州仁然文化传媒有限公司在杭州上城区注册，法人赵立业，注册资本50万元，参保0人。'
         '注册仅 30 天后 Medica 2024 开展。该公司的经营范围含"会议及展览服务"但主营业务为文化传媒/珠宝/服装——'
         '与医疗器械完全无关。这是一家专为签约参展而设的壳公司。',
         '证据: QCC 报告，成立日 2024-10-12; Medica 2024 展期 2024-11-11/14'),
        
        ('Step 8: 正式以新名入驻国际馆 (Medica 2025)',
         'AM MEDICAL EQUIPMENT PTE. LTD. 首次以新名在 Medica 2025 参展，展位从 Hall 6 D40 迁至 Hall 11 E69（国际馆核心区域）。'
         '同一展位上同时列出 Zhengzhou Olive Electronic Technology (CN) 作为共展商。'
         '两者均在数据库中标记为 high risk。Zhengzhou Olive 同时被标记为 cn_in_intl_hall (anomaly_flag #246)。',
         '证据: medica_investigation.db → participations + anomaly_flags'),
        
        ('Step 9: 拆售给下游中国企业 (持续)',
         '已确认 3 家下游中国企业通过 AM MEDICAL 展位参展:'
         '① 郑州奥利弗电子科技（11 E69, 2025）'
         '② 雷虎（河南）医疗器械（6 D40, 2025）'
         '③ 江苏康宝医疗器械（6 D40, 2025）'
         '拆售模式推断: 杭州仁然/赵立业通过深圳中检联的医疗器械检测客户资源寻找有出展需求的中国制造商 → '
         '签约后通过 AM MEDICAL 在德国直订国际馆展位 → 中国企业以"共展商"名义进驻。',
         '证据: medica_investigation.db 混合展位查询 + QCC 报告'),
    ]
    
    for title, desc, evidence in steps:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        doc.add_paragraph(desc)
        p2 = doc.add_paragraph()
        p2.add_run(evidence).italic = True
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ---- PART 4: 全球展会报名信息 ----
    add_heading_styled(doc, '第四部分 全球展会报名信息', 1)
    
    add_heading_styled(doc, '4.1 AM MEDICAL 全球参展记录', 2)
    
    s05 = load_json('findings/05_global_fairs.json')['findings']
    
    # Confirmed table
    doc.add_paragraph('已确认记录 (数据库验证):')
    table = doc.add_table(rows=len(s05['confirmed_in_db'])+1, cols=6, style='Light Grid Accent 1')
    headers = ['年份', '展会', '城市', '展位', '参展名义', '风险标签']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for i, r in enumerate(s05['confirmed_in_db']):
        cells = table.rows[i+1].cells
        cells[0].text = str(r['year'])
        cells[1].text = r['fair']
        cells[2].text = r['city']
        cells[3].text = r.get('booth', '-')
        cells[4].text = r['entity_name']
        cells[5].text = r.get('risk_label', '-')
    
    doc.add_paragraph()
    doc.add_paragraph('历史记录 (旧版SOP调研，待数据库补充验证):')
    
    table2 = doc.add_table(rows=len(s05['from_prior_sop_investigation'])+1, cols=5, style='Light Grid Accent 1')
    headers2 = ['年份', '展会', '城市', '展位', '参展名义']
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].text = h
    for i, r in enumerate(s05['from_prior_sop_investigation']):
        cells = table2.rows[i+1].cells
        cells[0].text = str(r['year'])
        cells[1].text = r['fair']
        cells[2].text = r['city']
        cells[3].text = r.get('booth', '未确认')
        cells[4].text = r['entity_name']
    
    doc.add_paragraph()
    doc.add_paragraph(f'模式分析: {s05["pattern_analysis"]["expansion_pattern"]}')
    doc.add_paragraph(f'地理覆盖: {", ".join(s05["pattern_analysis"]["geographic_span"])}')
    doc.add_paragraph(f'总计确认展会数: {s05["total_confirmed_fairs"]}')
    doc.add_paragraph(f'时间跨度: {s05["total_years_span"]}')
    
    add_heading_styled(doc, '4.2 已确认下游中国企业', 2)
    
    s06 = load_json('findings/06_downstream_cn.json')['findings']
    
    table3 = doc.add_table(rows=len(s06['confirmed_downstream'])+1, cols=5, style='Light Grid Accent 1')
    headers3 = ['企业名', '法人', '展位', '年份', '境外壳方']
    for i, h in enumerate(headers3):
        table3.rows[0].cells[i].text = h
    for i, r in enumerate(s06['confirmed_downstream']):
        cells = table3.rows[i+1].cells
        cells[0].text = r['name']
        cells[1].text = r.get('legal_rep', '待查')
        cells[2].text = r['booth']
        cells[3].text = str(r['year'])
        cells[4].text = r['shell_partner']
    
    doc.add_paragraph()
    doc.add_paragraph(f'总计已确认下游中国企业: {s06["total_confirmed_downstream"]} 家，在 {s06["total_confirmed_booths"]} 个展位。')
    doc.add_paragraph(f'注意: 此数字仅基于 medica_investigation.db (2025年数据)。AM MEDICAL 历年累计下游企业可能更多。')
    doc.add_paragraph(f'Medica 2025 全局背景: 数据库中存在大量同类"海外壳+中国企业"混合展位，AM MEDICAL 是其中一个可复现样本。')
    
    # ---- SUPPLEMENTARY ----
    doc.add_page_break()
    add_heading_styled(doc, '附录: 证据评级与方法论说明', 1)
    
    add_heading_styled(doc, 'A.1 证据强度定义', 2)
    doc.add_paragraph('HIGH (强证据): 可直接从官方数据库/API/WWW验证的事实，关联不可否认。如: 公司注册记录、WHOIS查询、数据库字段匹配、官方展商列表中的共展关系。')
    doc.add_paragraph('MED (中等证据): 来自可靠来源但有推断成分。如: 时间线巧合分析、行业模式匹配、经营范围异常信号。')
    doc.add_paragraph('LOW (弱证据/推断): 缺乏直接证据但逻辑上合理。需在报告中明确标注为"推断"。')
    
    add_heading_styled(doc, 'A.2 方法论审查', 2)
    doc.add_paragraph('• 所有数据库记录来自 medica_investigation.db (SQLite, 5674条Medica 2025参展记录 + 769条COMPAMED 2025记录)')
    doc.add_paragraph('• 所有中国工商信息来自 企查查 API (QCC) 已验证报告，raw JSON 已保存')
    doc.add_paragraph('• 所有新加坡工商信息来自 ACRA 公开记录 + companies.sg / sgbusinessdirectory.com')
    doc.add_paragraph('• WHOIS 信息来自 chinaz.com WHOIS 查询')
    doc.add_paragraph('• 报告中所有已证实关联均附来源URL或数据库字段路径')
    doc.add_paragraph('• 标注为"推断"的关联在报告中明确标识，不与已证实关联混淆')
    
    add_heading_styled(doc, 'A.3 局限性声明', 2)
    doc.add_paragraph('1. 新加坡 ACRA 官员姓名被 Cloudflare 拦截，无法在线获取董事/股东名单（需人工BizFile查询）')
    doc.add_paragraph('2. 数据库仅含 Medica 2025 + COMPAMED 2025 两届数据，2022-2024 年参展记录来自旧版SOP调研，待爬取补充')
    doc.add_paragraph('3. 赵立业与华尔集团/王超的直接关联（如股权、合同）不在工商公开信息中，无法通过公开渠道验证')
    doc.add_paragraph('4. 拆售金额估算基于 MEDICA 展位标准价格（~360 EUR/㎡），实际转售价格由买卖双方协商，无法获取')
    
    # Save
    outpath = os.path.join(BASE, 'outputs', f'AM_MEDICAL_背调报告_{today}.docx')
    doc.save(outpath)
    print(f'✅ 报告已生成: {outpath}')
    return outpath

if __name__ == '__main__':
    main()
