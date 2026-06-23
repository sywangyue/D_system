#!/usr/bin/env python3
"""生成 MWLAB-2026 系统汇报 Word 文档（4页，黑灰橙配色）"""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from pathlib import Path

# ── 颜色常量 ────────────────────────────────────────────────────────────────
ORANGE   = RGBColor(0xE8, 0x60, 0x2C)   # #E8602C  Messe 橙
DARK     = RGBColor(0x1A, 0x1A, 0x1A)   # #1A1A1A  近黑
MID_GRAY = RGBColor(0x55, 0x55, 0x55)   # #555555  正文灰
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99) # #999999  辅助灰
BG_ORANGE  = "E8602C"
BG_DARK    = "1A1A1A"
BG_LGRAY   = "F2F2F2"
BG_WHITE   = "FFFFFF"
BG_MGRAY   = "DDDDDD"

OUTPUT_PATH = Path(__file__).parent.parent / "MWLAB-2026_系统能力简报.docx"


def set_cell_bg(cell, hex_color: str):
    """设置单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, sides: list[str], color: str = "E8602C", sz: str = "12"):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), sz)
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_para(doc, text: str, size: int = 10, bold: bool = False,
             color: RGBColor = None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before: int = 0, space_after: int = 4) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or DARK


def add_run_mixed(para, segments: list[tuple]):
    """segments: [(text, bold, color, size)]"""
    for text, bold, color, size in segments:
        r = para.add_run(text)
        r.bold = bold
        r.font.color.rgb = color
        r.font.size = Pt(size)


def make_doc() -> Document:
    doc = Document()

    # ── 页面设置：A4，页边距 ──────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(1.6)
    section.bottom_margin = Cm(1.6)

    # ════════════════════════════════════════════════════════════════════════
    # PAGE 1 — 封面 + 系统概览
    # ════════════════════════════════════════════════════════════════════════

    # 封面标题块（深色背景表格）
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover.style = "Table Grid"
    cell = cover.cell(0, 0)
    set_cell_bg(cell, BG_DARK)
    cell.width = Cm(17)

    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cp.paragraph_format.space_before = Pt(14)
    cp.paragraph_format.space_after = Pt(2)
    r = cp.add_run("MWLAB-2026")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = ORANGE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run("展会竞争情报系统")
    r2.bold = True; r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(14)
    r3 = p3.add_run("系统能力简报  ·  2026年6月  ·  杜塞尔多夫展览（中国）")
    r3.font.size = Pt(9); r3.font.color.rgb = LIGHT_GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── 一句话价值主张 ──
    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    vp.paragraph_format.space_before = Pt(0)
    vp.paragraph_format.space_after = Pt(10)
    add_run_mixed(vp, [
        ("核心价值：", True, ORANGE, 11),
        ("将散落在市场上的展会数据，转化为BD团队可以直接使用的销售情报与竞争分析报告。", False, MID_GRAY, 10),
    ])

    # ── Section 1：数据从哪里来 ──
    _section_title(doc, "01", "数据从哪里来")

    src_table = doc.add_table(rows=2, cols=3)
    src_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    src_table.style = "Table Grid"

    src_data = [
        ("内部展会数据库\nmwlab.db",
         "收录中国市场\n5,941 个展会品牌\n6,129 届历史数据\n涵盖 8 大行业类别",
         BG_LGRAY),
        ("企查查\n企业工商数据",
         "实时查询企业注册\n信息、法定代表人、\n信用代码、经营状态\n（API接入）",
         BG_LGRAY),
        ("互联网公开信息\nWebSearch",
         "市场规模报告、\n行业趋势新闻、\n主办方背景、\n风险预警信息",
         BG_LGRAY),
    ]
    src_headers = ["数据源 ①", "数据源 ②", "数据源 ③"]

    # 表头行
    for j, hdr in enumerate(src_headers):
        c = src_table.cell(0, j)
        set_cell_bg(c, BG_ORANGE)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(hdr)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    # 内容行
    for j, (title, detail, bg) in enumerate(src_data):
        c = src_table.cell(1, j)
        set_cell_bg(c, bg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        rt = p.add_run(title.split("\n")[0] + "\n")
        rt.bold = True; rt.font.size = Pt(10); rt.font.color.rgb = ORANGE
        rd = p.add_run("\n".join(title.split("\n")[1:]) + "\n" + detail)
        rd.font.size = Pt(8.5); rd.font.color.rgb = MID_GRAY
        c.paragraphs[0].paragraph_format.space_after = Pt(6)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ════════════════════════════════════════════════════════════════════════
    # PAGE 2 — 4项核心能力
    # ════════════════════════════════════════════════════════════════════════

    _section_title(doc, "02", "系统能做什么（4项核心能力）")

    cap_data = [
        ("行业调研",
         "/industry-research",
         "某个行业里总共有多少展会？哪些规模最大？MDS 进入这个市场胜算几何？",
         ["全行业展会版图（数量/分布/规模）",
          "主要竞争对手盘点（主办方集中度）",
          "市场规模估算（TAM/SAM/SOM框架）",
          "波特五力分析（竞争强度评估）",
          "高优先级收购/合作目标清单"],
         "适用：开拓新行业前的市场可行性研判"),
        ("品牌展会调研",
         "/brand-research",
         "这个展会值不值得我们谈合作？规模在涨还是跌？主办方是谁？",
         ["历史届次规模趋势（近10年）",
          "竞争关系网络（同行业展会对比）",
          "主办方企业背景调查",
          "收并购价值评估（MA潜力评分）",
          "BD 谈判策略与时间窗口建议"],
         "适用：谈判前的标的尽调 & 竞品监控"),
        ("批量客户挖掘",
         "/batch-prospect",
         "竞品展会的参展商是我们的潜在客户，如何快速拿到一批名单？",
         ["输入竞品展会参展商关键词",
          "企查查批量搜索企业工商信息",
          "结果自动存库并导出 Excel 名单",
          "包含：公司名、法人、成立日期、状态、地址",
          "名单可直接用于BD电话/邮件开发"],
         "适用：竞品展会客户批量挖掘"),
        ("单一客户深度挖掘",
         "/single-prospect",
         "见这个客户前，我需要了解他的全部背景——包括有没有风险。",
         ["企业基本信息（工商 + 网络）",
          "展会参与历史（跨展会参展轨迹）",
          "潜在需求与接触切入点分析",
          "风险标注（违规代理商排查）",
          "BD行动建议（优先级+接触方式）"],
         "适用：重要客户拜访前的背景调研"),
    ]

    for i, (name, cmd, question, outputs, note) in enumerate(cap_data):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.style = "Table Grid"

        # 左列：标题 + 命令 + 问题
        lc = tbl.cell(0, 0)
        lc.width = Cm(6.5)
        set_cell_bg(lc, BG_DARK if i % 2 == 0 else "2D2D2D")

        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(8)
        lp.paragraph_format.space_after = Pt(2)
        r_num = lp.add_run(f"0{i+1}  ")
        r_num.font.size = Pt(14); r_num.bold = True; r_num.font.color.rgb = ORANGE
        r_name = lp.add_run(name)
        r_name.font.size = Pt(12); r_name.bold = True; r_name.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

        lp2 = lc.add_paragraph(cmd)
        lp2.paragraph_format.space_before = Pt(2)
        lp2.paragraph_format.space_after = Pt(6)
        lp2.runs[0].font.size = Pt(8.5)
        lp2.runs[0].font.color.rgb = LIGHT_GRAY

        lp3 = lc.add_paragraph(f'"{question}"')
        lp3.paragraph_format.space_before = Pt(2)
        lp3.paragraph_format.space_after = Pt(8)
        lp3.runs[0].font.size = Pt(8.5)
        lp3.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        lp3.runs[0].bold = True

        # 右列：输出清单 + 适用场景
        rc = tbl.cell(0, 1)
        rc.width = Cm(10.5)
        set_cell_bg(rc, BG_LGRAY)

        rp = rc.paragraphs[0]
        rp.paragraph_format.space_before = Pt(6)
        rp.paragraph_format.space_after = Pt(2)
        rr = rp.add_run("输出内容")
        rr.bold = True; rr.font.size = Pt(9); rr.font.color.rgb = ORANGE

        for item in outputs:
            ip = rc.add_paragraph()
            ip.paragraph_format.space_before = Pt(1)
            ip.paragraph_format.space_after = Pt(1)
            ip.paragraph_format.left_indent = Pt(8)
            ir = ip.add_run(f"▸  {item}")
            ir.font.size = Pt(8.5); ir.font.color.rgb = MID_GRAY

        np_ = rc.add_paragraph()
        np_.paragraph_format.space_before = Pt(5)
        np_.paragraph_format.space_after = Pt(6)
        nr = np_.add_run(note)
        nr.font.size = Pt(8); nr.font.color.rgb = LIGHT_GRAY; nr.italic = True

        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    # ════════════════════════════════════════════════════════════════════════
    # PAGE 3 — 项目成员视角 + 使用流程
    # ════════════════════════════════════════════════════════════════════════

    _section_title(doc, "03", "对项目组成员意味着什么")

    role_data = [
        ("BD经理 / 客户总监",
         [("过去", "靠经验和人脉判断市场，数据分散在Excel、PPT、个人脑子里"),
          ("现在", "30分钟内出一份行业报告 + 竞品分析，有数据支撑的谈判策略")]),
        ("客户开发 / 商务拓展",
         [("过去", "手工查企业官网、展商名录，效率低，遗漏多"),
          ("现在", "输入竞品展会关键词，自动生成企业名单Excel，直接打电话/发邮件")]),
        ("项目总监 / 运营",
         [("过去", "市场态势报告依赖主观判断，缺乏横向对比依据"),
          ("现在", "系统提供行业全量展会版图 + 量化评分，决策有据可查")]),
        ("新同事 / 实习生",
         [("过去", "不熟悉市场，不知道从哪入手做调研"),
          ("现在", "一条指令启动调研，系统自动汇总数据，产出专业水准报告")]),
    ]

    role_tbl = doc.add_table(rows=len(role_data), cols=3)
    role_tbl.style = "Table Grid"
    role_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (role, changes) in enumerate(role_data):
        rc0 = role_tbl.cell(i, 0)
        rc0.width = Cm(3.5)
        set_cell_bg(rc0, BG_ORANGE if i % 2 == 0 else "C24E1F")
        rp0 = rc0.paragraphs[0]
        rp0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rp0.paragraph_format.space_before = Pt(8)
        rp0.paragraph_format.space_after = Pt(8)
        rr0 = rp0.add_run(role.replace(" / ", "\n"))
        rr0.bold = True; rr0.font.size = Pt(9); rr0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

        for j, (label, desc) in enumerate(changes):
            col_idx = j + 1
            rc_j = role_tbl.cell(i, col_idx)
            rc_j.width = Cm(5.5) if j == 0 else Cm(6.5)
            set_cell_bg(rc_j, "EEEEEE" if label == "过去" else BG_WHITE)
            rp_j = rc_j.paragraphs[0]
            rp_j.paragraph_format.space_before = Pt(6)
            rp_j.paragraph_format.space_after = Pt(2)
            rl = rp_j.add_run(f"【{label}】  ")
            rl.bold = True
            rl.font.size = Pt(8.5)
            rl.font.color.rgb = LIGHT_GRAY if label == "过去" else ORANGE
            rd = rp_j.add_run(desc)
            rd.font.size = Pt(8.5)
            rd.font.color.rgb = MID_GRAY if label == "过去" else DARK
            rp_j.paragraph_format.space_after = Pt(6)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 操作流程示意 ──
    _section_title(doc, "04", "如何使用（三步启动一次调研）")

    flow_tbl = doc.add_table(rows=1, cols=5)
    flow_tbl.style = "Table Grid"
    flow_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    steps = [
        ("STEP 1", "打开终端\n输入指令", BG_ORANGE),
        ("→", "", "FFFFFF"),
        ("STEP 2", "系统自动\n抓取数据", BG_MGRAY),
        ("→", "", "FFFFFF"),
        ("STEP 3", "收到完整\nMarkdown报告\n（自动存档）", "1A1A1A"),
    ]
    for j, (title, desc, bg) in enumerate(steps):
        c = flow_tbl.cell(0, j)
        set_cell_bg(c, bg)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)

        if title == "→":
            r = p.add_run("→")
            r.bold = True; r.font.size = Pt(16)
            r.font.color.rgb = LIGHT_GRAY
            c.width = Cm(1.2)
        else:
            c.width = Cm(4)
            rt = p.add_run(title + "\n")
            rt.bold = True; rt.font.size = Pt(9)
            rt.font.color.rgb = RGBColor(0xFF,0xFF,0xFF) if bg in (BG_ORANGE,"1A1A1A") else DARK
            rd_p = c.add_paragraph(desc)
            rd_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rd_p.paragraph_format.space_before = Pt(0)
            rd_p.paragraph_format.space_after = Pt(8)
            rd_r = rd_p.runs[0]
            rd_r.font.size = Pt(8.5)
            rd_r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF) if bg in (BG_ORANGE,"1A1A1A") else MID_GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── 指令示例 ──
    ex_para = doc.add_paragraph()
    ex_para.paragraph_format.space_before = Pt(0)
    ex_para.paragraph_format.space_after = Pt(4)
    add_run_mixed(ex_para, [
        ("指令示例：  ", True, ORANGE, 9),
        ("/industry-research 机械和设备   ", False, MID_GRAY, 9),
        ("/brand-research 中国国际机床展览会   ", False, MID_GRAY, 9),
        ("/single-prospect 上海精密机床有限公司 代理商排查", False, MID_GRAY, 9),
    ])

    # ── 底部免责 ──
    footer_p = doc.add_paragraph()
    footer_p.paragraph_format.space_before = Pt(8)
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "系统当前覆盖中国境内展会数据，企查查工商查询需配置 API Key。"
        "所有报告自动存档，历史记录可追溯。"
        "  |  内部工具，不对外发布。"
    )
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = LIGHT_GRAY

    return doc


def _section_title(doc: Document, num: str, title: str):
    """橙色编号 + 黑色标题的分节标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_run_mixed(p, [
        (f"[ {num} ]  ", True, ORANGE, 11),
        (title, True, DARK, 11),
    ])
    # 下方橙色细线
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "E8602C")
    pBdr.append(bottom)
    pPr.append(pBdr)


if __name__ == "__main__":
    doc = make_doc()
    doc.save(str(OUTPUT_PATH))
    print(f"文档已生成：{OUTPUT_PATH}")
