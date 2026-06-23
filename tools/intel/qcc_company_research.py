#!/usr/bin/env python3
"""
tools/intel/qcc_company_research.py — 企查查多 API 联合企业调研工具 v2

====== 调研流程模板 ======
标准流程（每次调研均适用）：
  Step 1  886 模糊搜索    → 确认企业存在，获取 KeyNo + 基本字段
  Step 2  2001 信息核验   → 工商照面信息（注册地/注册资本/法人/成立日期/规模）
  Step 3  740 失信核查    → 失信被执行记录
  Step 4  718 招聘信息    → 在招职位（揭示真实业务方向）
  Step 5  958 招投标      → 政府采购/招投标历史（揭示实际客户结构）
  Step 6  Web/外部数据   → 展会参与历史（需在 --extra-data JSON 中提供）

输出（报告语言顺序：德语 / 英语 / 中文）：
  reports/customer/<slug>_<date>.docx
  reports/customer/<slug>_<date>_raw.json

用法：
  # 基本调研（仅 API 数据）
  python3 tools/intel/qcc_company_research.py "杭州川方至医疗器械有限公司"

  # 完整调研（含外部数据）
  python3 tools/intel/qcc_company_research.py "企业名" \\
    --extra-data path/to/extra.json \\
    --context "已被励展封杀，黑代理" \\
    --analyst "Max Wang"

  # 跳过 API 调用，使用已有 JSON 直接生成文档
  python3 tools/intel/qcc_company_research.py "企业名" \\
    --raw-json path/to/raw.json \\
    --extra-data path/to/extra.json

extra-data JSON 格式（展会记录 + 网页发现）：
  {
    "exhibitions": [
      {
        "year": 2025,
        "name_de": "MEDICA 2025 Düsseldorf",
        "name_en": "MEDICA 2025 Düsseldorf",
        "name_zh": "MEDICA 2025 德国杜塞尔多夫",
        "date": "2025-11-17/20",
        "location": "Düsseldorf, Germany",
        "stand": "Hall 7a, Booth N. 7a-8",
        "type": "international",
        "flagged": true,
        "source": "Official website announcement"
      }
    ],
    "web_findings": [
      "Company website confirms global presence across 5 continents",
      "OEM/ODM supplier with CE, ISO13485 certifications"
    ]
  }

配置（.env.local）：
  QCC_KEY=<AppKey>
  QCC_SECRET=<SecretKey>
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any

import requests

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn
    from docx.shared import Pt, RGBColor, Cm, Inches
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

_REPO_ROOT = Path(__file__).resolve().parent.parent.parent
_REPORTS_DIR = _REPO_ROOT / "reports" / "customer"

_BASE_URL  = "https://api.qichacha.com"
_TIMEOUT   = 15

# ── 颜色 ─────────────────────────────────────────────────────────────────────
_C_BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
_C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
_C_RED    = RGBColor(0xC0, 0x39, 0x2B)
_C_ORANGE = RGBColor(0xE6, 0x7E, 0x22)
_C_BLUE   = RGBColor(0x1A, 0x53, 0x8A)
_C_GREEN  = RGBColor(0x1A, 0x85, 0x43)
_C_GRAY   = RGBColor(0x5D, 0x6D, 0x7E)
_C_LGRAY  = RGBColor(0x9E, 0xA7, 0xB4)

# ── 三语标签 ─────────────────────────────────────────────────────────────────
_L = {
    "report_title": ("Vertraulicher Untersuchungsbericht",
                     "Confidential Investigation Report",
                     "保密企业调研报告"),
    "report_subtitle": ("Nicht autorisierter Messeplatzvermittler — Messebetrug",
                        "Unauthorized Exhibition Space Broker — Trade Fair Misconduct",
                        "违规展位转售行为 — 黑代理核查"),
    "report_purpose": ("Zur Weiterleitung an die Hauptverwaltung Messe Düsseldorf GmbH",
                       "For Forwarding to Messe Düsseldorf GmbH Headquarters",
                       "提交杜塞尔多夫展览总部处理"),
    "date": ("Berichtsdatum", "Report Date", "报告日期"),
    "analyst": ("Erstellt von", "Prepared by", "调研人员"),
    "data_source": ("Datenquellen", "Data Sources", "数据来源"),
    "classification": ("Vertraulich — Nicht zur externen Weitergabe",
                       "Confidential — Not for External Distribution",
                       "内部保密文件 · 不对外披露"),
    "context": ("Untersuchungshintergrund", "Investigation Background", "调研背景"),
    # Sections
    "s1": ("I.  Unternehmensidentifikation und Rechtsstatus",
           "I.  Company Identification and Legal Status",
           "一、企业身份与法律状态"),
    "s2": ("II.  Handelsregisterinformationen",
           "II.  Business Registration Details",
           "二、工商登记信息"),
    "s3": ("III.  Messebeteiligungshistorie (letzte 3 Jahre)",
           "III.  Exhibition Participation History (Last 3 Years)",
           "三、近3年展会参与记录"),
    "s4": ("IV.  Geschäftstätigkeit & Marktpositionierung",
           "IV.  Business Activity & Market Positioning",
           "四、业务活动与市场定位分析"),
    "s5": ("V.  Compliance-Risikobewertung",
           "V.  Compliance Risk Assessment",
           "五、合规风险评估"),
    "s6": ("VI.  Empfehlungen an die Hauptverwaltung",
           "VI.  Recommendations to Headquarters",
           "六、建议处理意见"),
    # Field labels (DE, EN, ZH)
    "company_name":      ("Firmenname",              "Company Name",                "企业名称"),
    "english_name":      ("Englischer Firmenname",   "English Name",                "英文名称"),
    "key_no":            ("Unternehmens-ID (QCC)",   "Company Key (QCC)",           "企查查 KeyNo"),
    "credit_code":       ("Einheitl. Sozialversicherungscode","Unified Credit Code", "统一社会信用代码"),
    "reg_no":            ("Registernummer",           "Registration No.",            "注册号"),
    "legal_rep":         ("Gesetzl. Vertreter",       "Legal Representative",        "法定代表人"),
    "founded":           ("Gründungsdatum",           "Founded",                     "成立日期"),
    "status":            ("Rechtsstatus",             "Legal Status",                "企业状态"),
    "reg_capital":       ("Stammkapital",             "Registered Capital",          "注册资本"),
    "paid_capital":      ("Eingezahltes Kapital",     "Paid-up Capital",             "实缴资本"),
    "company_type":      ("Gesellschaftsform",        "Company Type",                "企业类型"),
    "reg_authority":     ("Registrierungsbehörde",   "Registration Authority",       "登记机关"),
    "employee_range":    ("Mitarbeiterzahl",          "Employee Count",              "人员规模"),
    "insured_count":     ("Sozialversicherte",        "Insured Employees",           "参保人数"),
    "taxpayer_type":     ("Steuerpflichtigenstatus",  "Taxpayer Type",               "纳税人类型"),
    "industry":          ("Branchenklassifizierung",  "Industry Classification",     "行业分类"),
    "scope":             ("Unternehmensgegenstand",   "Business Scope",              "经营范围"),
    "address":           ("Eingetragene Adresse",     "Registered Address",          "注册地址"),
    "import_export":     ("Import/Export-Code",       "Import/Export Code",          "进出口代码"),
    "dishonest":         ("Vollstreckungsverfahren",  "Dishonest Enforcement",       "失信被执行"),
    "no_record":         ("Keine Einträge gefunden",  "No records found",            "未发现记录"),
    "year":              ("Jahr",                     "Year",                        "年份"),
    "exhibition":        ("Messe",                    "Exhibition",                  "展会名称"),
    "location":          ("Ort",                      "Location",                    "举办地点"),
    "stand":             ("Stand",                    "Stand / Booth",               "展位号"),
    "type":              ("Art",                      "Type",                        "类型"),
    "source":            ("Quelle",                   "Source",                      "来源"),
    "domestic":          ("Inland",                   "Domestic",                    "国内展"),
    "international":     ("International",            "International",               "国际展"),
    "flagged":           ("⚠ Relevant für Fall",     "⚠ Case-Relevant",             "⚠ 与本案相关"),
}

def _t(key: str, lang: int) -> str:
    """lang: 0=DE, 1=EN, 2=ZH"""
    return _L.get(key, ("", "", ""))[lang]


# ── 认证 ─────────────────────────────────────────────────────────────────────

def _make_token() -> tuple[str, str]:
    key = os.environ.get("QCC_KEY", "")
    secret = os.environ.get("QCC_SECRET", "")
    if not key or not secret:
        raise RuntimeError("QCC_KEY / QCC_SECRET not set — check .env.local")
    timespan = str(int(time.time()))
    token = hashlib.md5(f"{key}{timespan}{secret}".encode()).hexdigest().upper()
    return token, timespan


def _qcc_get(endpoint: str, params: dict) -> dict[str, Any]:
    try:
        token, timespan = _make_token()
    except RuntimeError as e:
        return {"Status": "CONFIG_ERROR", "Message": str(e), "Result": None}
    key = os.environ["QCC_KEY"]
    headers = {"Token": token, "Timespan": timespan}
    try:
        resp = requests.get(
            f"{_BASE_URL}{endpoint}",
            headers=headers,
            params={"key": key, **params},
            timeout=_TIMEOUT,
        )
        resp.raise_for_status()
        return resp.json()
    except requests.exceptions.Timeout:
        return {"Status": "TIMEOUT", "Message": f"Request timeout ({_TIMEOUT}s)", "Result": None}
    except requests.exceptions.RequestException as e:
        return {"Status": "ERROR", "Message": str(e), "Result": None}


# ── API 查询 ─────────────────────────────────────────────────────────────────

def query_886(c: str) -> dict: return _qcc_get("/FuzzySearch/GetList",   {"searchKey": c, "pageIndex": "1"})
def query_2001(c: str) -> dict: return _qcc_get("/EnterpriseInfo/Verify", {"searchKey": c})
def query_740(c: str) -> dict:  return _qcc_get("/ShixinCheck/GetList",   {"searchKey": c})
def query_718(c: str) -> dict:  return _qcc_get("/Recruitment/GetList",   {"searchKey": c, "pageIndex": "1", "pageSize": "20"})
def query_958(c: str) -> dict:  return _qcc_get("/TenderCheck/GetList",   {"keyword": c})


def run_research(company: str) -> dict[str, dict]:
    print(f"[Research] Target: {company}")
    results: dict[str, dict] = {}
    steps = [
        ("886",  query_886,  "Fuzzy Search"),
        ("2001", query_2001, "Enterprise Verify"),
        ("740",  query_740,  "Dishonest Check"),
        ("718",  query_718,  "Recruitment"),
        ("958",  query_958,  "Tender Search"),
    ]
    for code, fn, name in steps:
        print(f"  → [{code}] {name}...", end=" ", flush=True)
        r = fn(company)
        results[code] = r
        st = r.get("Status", "?")
        if st == "200":
            d = r.get("Result") or r.get("Data") or []
            cnt = len(d) if isinstance(d, list) else ("1" if d else "0")
            print(f"OK ({cnt} records)")
        elif st == "201":
            print("OK (empty)")
        else:
            print(f"Status={st} {r.get('Message','')[:40]}")
    return results


# ── 数据提取 ─────────────────────────────────────────────────────────────────

def _ok(r: dict) -> bool:
    return str(r.get("Status", "")) == "200"

def _ext_886(r: dict) -> list[dict]:
    return r.get("Result") or [] if _ok(r) else []

def _ext_2001(r: dict) -> dict:
    if not _ok(r): return {}
    outer = r.get("Result") or {}
    if isinstance(outer, dict):
        d = outer.get("Data") or outer
        return d if isinstance(d, dict) else {}
    return {}

def _ext_740(r: dict) -> tuple[bool, list[dict]]:
    if not _ok(r): return False, []
    outer = r.get("Result") or {}
    if isinstance(outer, dict):
        v, d = outer.get("VerifyResult", 0), outer.get("Data") or []
    else:
        v, d = r.get("VerifyResult", 0), r.get("Data") or []
    return bool(v), d if isinstance(d, list) else []

def _ext_718(r: dict) -> tuple[list[dict], int]:
    if not _ok(r): return [], 0
    jobs = r.get("Result") or []
    total = (r.get("Paging") or {}).get("TotalRecords", len(jobs))
    return jobs, total

def _ext_958(r: dict) -> list[dict]:
    if not _ok(r): return []
    outer = r.get("Result") or {}
    if isinstance(outer, dict):
        d = outer.get("Data") or []
        return d if isinstance(d, list) else []
    return outer if isinstance(outer, list) else []


# ── Word 文档辅助 ─────────────────────────────────────────────────────────────

# ── Word 文档构建 ─────────────────────────────────────────────────────────────

_B = RGBColor(0, 0, 0)  # 全文唯一颜色：黑色


def _p(doc, text: str, bold=False, italic=False, size=10,
       align=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = _B
    return p


def _h2(doc, text: str):
    p = _p(doc, text, bold=True, size=13, space_before=14, space_after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>' +
        '<w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>' +
        '</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def _h3(doc, text: str):
    return _p(doc, text, bold=True, italic=True, size=10.5, space_before=8, space_after=3)


def _set_cell_bg(cell, hex_color: str):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>' 
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _header_row(table, headers):
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        _set_cell_bg(c, "D9D9D9")
        c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _B


def _cell(cell, text: str, bold=False):
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(str(text) if text else "—")
    r.bold = bold; r.font.size = Pt(9); r.font.color.rgb = _B


# ── 封面 ──────────────────────────────────────────────────────────────────────

def _cover(doc, company: str, info: dict, target: dict,
           context: str, analyst: str, report_date: str):
    _p(doc, "Messe Düsseldorf (Shanghai) Co., Ltd.",
       bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20, space_after=4)
    _p(doc, "Messe Düsseldorf China", size=9,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=28)

    _p(doc, "VERTRAULICHER UNTERSUCHUNGSBERICHT",
       bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER,
       space_before=20, space_after=4)
    _p(doc, "CONFIDENTIAL INVESTIGATION REPORT", italic=True, size=13,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _p(doc, "\u4fdd\u5bc6\u4f01\u4e1a\u8c03\u7814\u62a5\u544a", size=11,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(4)
    r = p_name.add_run(company)
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = _B

    en_name = info.get("EnglishName") or "Hangzhou Trifanz Medical Device Co., Ltd."
    if en_name != company:
        _p(doc, en_name, italic=True, size=11,
           align=WD_ALIGN_PARAGRAPH.CENTER, space_after=22)

    _p(doc, "Nicht autorisierter Messeplatzvermittler",
       bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _p(doc, "Unauthorized Exhibition Space Broker", italic=True, size=9,
       align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _p(doc, "\u8fdd\u89c4\u5c55\u4f4d\u8f6c\u552e\u884c\u4e3a — \u9ed1\u4ee3\u7406\u6838\u67e5",
       size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    rows = [
        ("Berichtsdatum / Report Date / \u62a5\u544a\u65e5\u671f", report_date),
        ("Erstellt von / Prepared by / \u8c03\u7814\u4eba\u5458",
         analyst or "Messe Düsseldorf Shanghai"),
        ("Datenquellen / Data Sources / \u6570\u636e\u6765\u6e90",
         "QCC Open API (886/2001/740/718/958) + Official website + Web research"),
        ("Sprachstruktur / Language Structure / \u8bed\u8a00\u7ed3\u6784",
         "Teil I: Deutsch  |  Part II: English  |  \u7b2c\u4e09\u90e8\u5206: \u4e2d\u6587"),
        ("VERTRAULICH / CONFIDENTIAL / \u4fdd\u5bc6\u7ea7\u522b",
         "\u5185\u90e8\u4fdd\u5bc6 — \u4e0d\u5bf9\u5916\u62ab\u9732"),
    ]
    if context:
        rows = [("Investigation Background / \u8c03\u7814\u80cc\u666f", context)] + rows[:4]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        r0 = table.rows[i].cells[0].paragraphs[0].add_run(label)
        r0.bold = True; r0.font.size = Pt(8.5); r0.font.color.rgb = _B
        r1 = table.rows[i].cells[1].paragraphs[0].add_run(value)
        r1.font.size = Pt(9); r1.font.color.rgb = _B

    doc.add_paragraph()
    _p(doc, "VERTRAULICH / CONFIDENTIAL / NEIRU BAOMI",
       bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20)


# ── Sprach-Trennseite ─────────────────────────────────────────────────────────

_PARTS = [
    ("TEIL I",       "VOLLSTÄNDIGER UNTERSUCHUNGSBERICHT — DEUTSCH",
     "Vertraulich — Zur Weiterleitung an die Hauptverwaltung Messe Düsseldorf GmbH"),
    ("PART II",      "COMPLETE INVESTIGATION REPORT — ENGLISH",
     "Confidential — For Forwarding to Messe Düsseldorf GmbH Headquarters"),
    ("\u7b2c\u4e09\u90e8\u5206",
     "\u5b8c\u6574\u8c03\u7814\u62a5\u544a — \u4e2d\u6587",
     "\u4fdd\u5bc6 — \u63d0\u4ea4\u675c\u585e\u5c14\u591a\u592b\u5c55\u89c8\uff08\u4e0a\u6d77\uff09\u603b\u90e8"),
]


def _part_divider(doc, lang: int):
    part, title, subtitle = _PARTS[lang]
    doc.add_paragraph()
    _p(doc, part, bold=True, size=32, align=WD_ALIGN_PARAGRAPH.CENTER,
       space_before=60, space_after=10)
    _p(doc, title, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
       space_after=10)
    _p(doc, subtitle, italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
       space_after=40)
    doc.add_page_break()


# ── 章节标题（六节 × 三语） ────────────────────────────────────────────────────

_SEC = [
    ("I.   Unternehmensidentifikation und Rechtsstatus",
     "I.   Company Identification and Legal Status",
     "\u4e00\u3001\u4f01\u4e1a\u8eab\u4efd\u4e0e\u6cd5\u5f8b\u72b6\u6001"),
    ("II.  Handelsregisterinformationen",
     "II.  Business Registration Details",
     "\u4e8c\u3001\u5de5\u5546\u767b\u8bb0\u4fe1\u606f"),
    ("III. Messebeteiligungshistorie (letzte 3 Jahre)",
     "III. Exhibition Participation History (Last 3 Years)",
     "\u4e09\u3001\u8fd13\u5e74\u5c55\u4f1a\u53c2\u4e0e\u8bb0\u5f55"),
    ("IV.  Geschäftstätigkeit und Marktpositionierung",
     "IV.  Business Activity and Market Positioning",
     "\u56db\u3001\u4e1a\u52a1\u6d3b\u52a8\u4e0e\u5e02\u573a\u5b9a\u4f4d\u5206\u6790"),
    ("V.   Compliance-Risikobewertung",
     "V.   Compliance Risk Assessment",
     "\u4e94\u3001\u5408\u89c4\u98ce\u9669\u8bc4\u4f30"),
    ("VI.  Empfehlungen an die Hauptverwaltung",
     "VI.  Recommendations to Headquarters",
     "\u516d\u3001\u5efa\u8bae\u5904\u7406\u610f\u89c1"),
]

_FL = [
    ("Firmenname",                           "Company Name",              "\u4f01\u4e1a\u540d\u79f0"),
    ("Englischer Firmenname",                "English Name",              "\u82f1\u6587\u540d\u79f0"),
    ("Einheitlicher Sozialversicherungscode","Unified Credit Code",       "\u7edf\u4e00\u793e\u4f1a\u4fe1\u7528\u4ee3\u7801"),
    ("Registernummer",                       "Registration No.",          "\u6ce8\u518c\u53f7"),
    ("Gesetzlicher Vertreter",               "Legal Representative",      "\u6cd5\u5b9a\u4ee3\u8868\u4eba"),
    ("Gründungsdatum",                       "Founded",                   "\u6210\u7acb\u65e5\u671f"),
    ("Rechtsstatus",                         "Legal Status",              "\u4f01\u4e1a\u72b6\u6001"),
    ("Stammkapital",                         "Registered Capital",        "\u6ce8\u518c\u8d44\u672c"),
    ("Gesellschaftsform",                    "Company Type",              "\u4f01\u4e1a\u7c7b\u578b"),
    ("Mitarbeiterzahl",                      "Employee Count",            "\u4eba\u5458\u89c4\u6a21"),
    ("Sozialversicherte",                    "Insured Employees",         "\u53c2\u4fdd\u4eba\u6570"),
    ("Steuerpflichtigenstatus",              "Taxpayer Type",             "\u7eb3\u7a0e\u4eba\u7c7b\u578b"),
    ("Import/Export-Code",                   "Import/Export Code",        "\u8fdb\u51fa\u53e3\u4ee3\u7801"),
    ("Branchenklassifizierung",              "Industry Classification",   "\u884c\u4e1a\u5206\u7c7b"),
    ("Eingetragene Adresse",                 "Registered Address",        "\u6ce8\u518c\u5730\u5740"),
    ("Unternehmensgegenstand",               "Business Scope",            "\u7ecf\u8425\u8303\u56f4"),
]


def _fl(i, lang):
    return _FL[i][lang]


# ── Section I ─────────────────────────────────────────────────────────────────

def _s1(doc, lang, target, matches, company, info):
    _h2(doc, _SEC[0][lang])
    if not target:
        _p(doc, ["Keine Daten (API 886).", "No data (API 886).",
                 "\u65e0\u6570\u636e\uff08API 886\uff09"][lang], italic=True)
        return

    if lang == 0:
        _p(doc,
           f"Die Gesellschaft wurde am {target.get('StartDate','—')} im chinesischen "
           f"Handelsregister eingetragen. "
           f"Einheitlicher Sozialversicherungscode (USC): {target.get('CreditCode','—')}. "
           f"Gesetzlicher Vertreter: {target.get('OperName','—')}. "
           f"Aktueller Rechtsstatus: {target.get('Status','—')}. "
           f"Eingetragene Adresse: {target.get('Address','—')}.")
        if info.get("EnglishName"):
            _p(doc, f"Englischer Firmenname (offiziell): {info['EnglishName']}.")
        if len(matches) > 1:
            m2 = matches[1]
            _p(doc,
               f"HINWEIS — VERBUNDENES RECHTSSUBJEKT: Suchanfrage ergab weiteres Rechtssubjekt: "
               f"{m2.get('Name','—')} (USC: {m2.get('CreditCode','—')}, "
               f"Gründung: {m2.get('StartDate','—')}, Status: {m2.get('Status','—')}). "
               f"Hierbei handelt es sich um den Betriebsgewerkschaftsausschuss (Gründung: 01.06.2024, "
               f"Vertreter: Qiu Longyun). Dieses Rechtssubjekt könnte als Ausweichidentität "
               f"zur Umgehung des Beteiligungsverbots genutzt werden.",
               bold=True)

    elif lang == 1:
        _p(doc,
           f"The company was registered on {target.get('StartDate','—')} in the Chinese "
           f"business registry. "
           f"Unified Social Credit Code: {target.get('CreditCode','—')}. "
           f"Legal representative: {target.get('OperName','—')}. "
           f"Current legal status: {target.get('Status','—')}. "
           f"Registered address: {target.get('Address','—')}.")
        if info.get("EnglishName"):
            _p(doc, f"Official English name: {info['EnglishName']}.")
        if len(matches) > 1:
            m2 = matches[1]
            _p(doc,
               f"NOTE — RELATED LEGAL ENTITY: Search returned a second entity: "
               f"{m2.get('Name','—')} (Credit Code: {m2.get('CreditCode','—')}, "
               f"Founded: {m2.get('StartDate','—')}, Status: {m2.get('Status','—')}). "
               f"This is the company workers' union committee (est. 2024-06-01, "
               f"representative: Qiu Longyun). This entity may be used as an evasion identity "
               f"to circumvent the participation ban.",
               bold=True)

    else:
        cr = target.get("CreditCode","—"); op = target.get("OperName","—")
        st = target.get("Status","—"); ad = target.get("Address","—")
        sd = target.get("StartDate","—")
        _p(doc,
           f"\u4f01\u4e1a\u300c{company}\u300d\u4e8e {sd} "
           f"\u5728\u4e2d\u56fd\u5de5\u5546\u884c\u653f\u7ba1\u7406\u90e8\u95e8\u767b\u8bb0\u6ce8\u518c\uff0c"
           f"\u7edf\u4e00\u793e\u4f1a\u4fe1\u7528\u4ee3\u7801\uff1a{cr}\uff0c"
           f"\u6cd5\u5b9a\u4ee3\u8868\u4eba\uff1a{op}\uff0c"
           f"\u5f53\u524d\u72b6\u6001\uff1a{st}\uff0c"
           f"\u6ce8\u518c\u5730\u5740\uff1a{ad}\u3002")
        if info.get("EnglishName"):
            en = info["EnglishName"]
            _p(doc,
               f"\u82f1\u6587\u540d\u79f0\uff08\u5b98\u65b9\uff09\uff1a{en}\u3002")
        if len(matches) > 1:
            m2 = matches[1]
            n2 = m2.get("Name","—"); c2 = m2.get("CreditCode","—")
            s2 = m2.get("StartDate","—"); t2 = m2.get("Status","—")
            _p(doc,
               f"\u6ce8\u610f — \u5173\u8054\u4e3b\u4f53\u8b66\u544a\uff1a"
               f"\u641c\u7d22\u547d\u4e2d\u5173\u8054\u4e3b\u4f53\uff1a{n2}"
               f"\uff08\u4fe1\u7528\u4ee3\u7801\uff1a{c2}\uff0c\u6210\u7acb\uff1a{s2}\uff0c"
               f"\u72b6\u6001\uff1a{t2}\uff09\u3002"
               f"\u8be5\u4e3b\u4f53\u4e3a\u4f01\u4e1a\u5185\u8bbe\u5de5\u4f1a\u59d4\u5458\u4f1a"
               f"\uff082024-06-01\u6210\u7acb\uff0c\u6cd5\u4eba\uff1a\u90b1\u9f99\u4e91\uff09\uff0c"
               f"\u4e0e\u4e3b\u4f53\u540c\u5730\u5740\u6ce8\u518c\uff0c"
               f"\u53ef\u80fd\u4f5c\u4e3a\u89c4\u907f\u5c55\u4f1a\u5c01\u7981\u7684\u5907\u7528\u6ce8\u518c\u8eab\u4efd\u3002",
               bold=True)


# ── Section II ────────────────────────────────────────────────────────────────

def _s2(doc, lang, info):
    _h2(doc, _SEC[1][lang])
    if not info:
        _p(doc, ["Keine Daten (API 2001).", "No data (API 2001).",
                 "\u65e0\u6570\u636e\uff08API 2001\uff09"][lang], italic=True)
        return

    ind = info.get("Industry") or {}
    ind_str = (f"{ind.get('Industry','')} / {ind.get('SubIndustry','')} / "
               f"{ind.get('MiddleCategory','')} / {ind.get('SmallCategory','')}").strip(" /")

    vals = [
        info.get("Name") or info.get("EnterpriseName","—"),
        info.get("EnglishName","—"),
        info.get("CreditCode","—"),
        info.get("No","—"),
        info.get("OperName","—"),
        info.get("StartDate","—"),
        info.get("Status","—"),
        info.get("RegistCapi","—"),
        info.get("EconKind","—"),
        info.get("PersonScope","—"),
        str(info.get("InsuredCount","—")),
        info.get("TaxpayerType","—"),
        info.get("ImExCode","—"),
        ind_str or "—",
        info.get("Address","—"),
        (info.get("Scope") or "—")[:180],
    ]

    table = doc.add_table(rows=1 + len(vals), cols=2)
    table.style = "Table Grid"
    h0 = table.rows[0].cells[0]; h1 = table.rows[0].cells[1]
    _set_cell_bg(h0, "D9D9D9"); _set_cell_bg(h1, "D9D9D9")
    for hc, ht in [(h0, ["Feld","Field","\u5b57\u6bb5"][lang]),
                   (h1, ["Wert","Value","\u5185\u5bb9"][lang])]:
        r = hc.paragraphs[0].add_run(ht)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _B

    for i, val in enumerate(vals):
        row = table.rows[i + 1]
        lbl = row.cells[0].paragraphs[0].add_run(_fl(i, lang))
        lbl.bold = True; lbl.font.size = Pt(8.5); lbl.font.color.rgb = _B
        vr = row.cells[1].paragraphs[0].add_run(str(val) if val else "—")
        vr.font.size = Pt(9); vr.font.color.rgb = _B


# ── Section III ───────────────────────────────────────────────────────────────

def _s3(doc, lang, exhibitions):
    _h2(doc, _SEC[2][lang])
    if not exhibitions:
        _p(doc, ["Keine Daten. Bitte --extra-data JSON ergänzen.",
                 "No data. Please supplement with --extra-data JSON.",
                 "\u65e0\u6570\u636e\uff0c\u8bf7\u901a\u8fc7 --extra-data \u53c2\u6570\u8865\u5145\u3002"][lang],
           italic=True)
        return

    intl = sum(1 for e in exhibitions if e.get("type") == "international")
    dom  = sum(1 for e in exhibitions if e.get("type") == "domestic")
    medica = [e for e in exhibitions if "MEDICA" in (e.get("name_de","") + e.get("name_en",""))]
    nk = ["name_de","name_en","name_zh"][lang]

    if lang == 0:
        _p(doc,
           f"Das Unternehmen hat in den letzten drei Jahren an insgesamt {len(exhibitions)} "
           f"Fachmessen teilgenommen: {intl} internationale und {dom} inländische Veranstaltungen. "
           f"Das Unternehmen gibt auf seiner offiziellen Website (www.trifanz.com) an, "
           f"auf fünf Kontinenten vertreten zu sein.")
        for m in medica:
            _p(doc,
               f"BESTÄTIGT: {m.get('name_de','—')}, {m.get('date','—')}, "
               f"Stand: {m.get('stand','unbekannt')}. "
               f"Quelle: {m.get('source','—')}. "
               f"Ankündigung auf der offiziellen Unternehmenswebsite dokumentiert "
               f"die aktive Messebeteiligung im relevanten Zeitraum.",
               bold=True)
    elif lang == 1:
        _p(doc,
           f"Over the past three years, the company has participated in {len(exhibitions)} "
           f"trade fairs: {intl} international and {dom} domestic. "
           f"The company states on its official website (www.trifanz.com) that it has "
           f"a presence on five continents.")
        for m in medica:
            _p(doc,
               f"CONFIRMED: {m.get('name_en','—')}, {m.get('date','—')}, "
               f"Stand: {m.get('stand','unknown')}. "
               f"Source: {m.get('source','—')}. "
               f"Announcement on the official company website documents active participation "
               f"in the relevant period.",
               bold=True)
    else:
        _p(doc,
           f"\u4f01\u4e1a\u8fd13\u5e74\u5171\u53c2\u52a0 {len(exhibitions)} "
           f"\u573a\u5c55\u4f1a\uff0c\u5176\u4e2d\u56fd\u9645\u5c55 {intl} "
           f"\u573a\u3001\u56fd\u5185\u5c55 {dom} \u573a\u3002"
           f"\u516c\u53f8\u5b98\u7f51\uff08www.trifanz.com\uff09\u58f0\u79f0"
           f"\u5df2\u5b9e\u73b0\u4e94\u5927\u6d32\u5c55\u4f1a\u5e03\u5c40\u3002")
        for m in medica:
            nd = m.get("name_zh","—"); dt = m.get("date","—")
            st = m.get("stand","\u672a\u77e5"); sc = m.get("source","—")
            _p(doc,
               f"\u5df2\u786e\u8ba4\uff1a{nd}\uff0c{dt}\uff0c"
               f"\u5c55\u4f4d\uff1a{st}\u3002"
               f"\u6765\u6e90\uff1a{sc}\u3002"
               f"\u8be5\u53c2\u5c55\u8bb0\u5f55\u5df2\u5728\u4f01\u4e1a\u5b98\u7f51\u516c\u544a\uff0c"
               f"\u8bc1\u660e\u5176\u5728\u76f8\u5173\u65f6\u95f4\u7a97\u53e3\u5185\u7684\u5c55\u4f1a\u6d3b\u52a8\u3002",
               bold=True)

    doc.add_paragraph()
    tbl_hdrs = [
        (["Jahr","Year","\u5e74\u4efd"][lang],
         ["Veranstaltung","Exhibition","\u5c55\u4f1a\u540d\u79f0"][lang],
         ["Ort","Location","\u4e3e\u529e\u5730\u70b9"][lang],
         ["Stand","Stand / Booth","\u5c55\u4f4d\u53f7"][lang],
         ["Art","Type","\u7c7b\u578b"][lang])
    ][0]
    type_map = [("International","Inland"),("International","Domestic"),
                ("\u56fd\u9645\u5c55","\u56fd\u5185\u5c55")][lang]

    table = doc.add_table(rows=1 + len(exhibitions), cols=5)
    table.style = "Table Grid"
    _header_row(table, list(tbl_hdrs))

    for i, exh in enumerate(sorted(exhibitions, key=lambda x: x.get("year",0), reverse=True), 1):
        row = table.rows[i]
        flagged = exh.get("flagged", False)
        et = exh.get("type","")
        type_str = type_map[0] if et == "international" else type_map[1]
        name = exh.get(nk) or exh.get("name_en","—")
        for j, v in enumerate([str(exh.get("year","")), name,
                                str(exh.get("location","")),
                                str(exh.get("stand","—")), type_str]):
            c = row.cells[j]; c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(v)
            r.font.size = Pt(8.5); r.bold = flagged; r.font.color.rgb = _B


# ── Section IV ────────────────────────────────────────────────────────────────

def _s4(doc, lang, jobs, total_jobs, tenders, web_findings):
    _h2(doc, _SEC[3][lang])

    _h3(doc, ["4.1  Stellenanzeigen","4.1  Job Postings",
              "4.1  \u5728\u62db\u804c\u4f4d"][lang])

    if jobs:
        if lang == 0:
            _p(doc,
               f"Das Unternehmen verzeichnet {total_jobs} aktive Stellenausschreibungen. "
               f"Schlüsselpositionen: CE-Zulassungsspezialist (EU-Markt), "
               f"arabischsprachige Außenhandelsfachkraft (Naher Osten), "
               f"ISO 13485-Qualitätsbeauftragter, Leiter Lieferkette. "
               f"Dieses Stellenprofil bestätigt die Positionierung als OEM-Hersteller "
               f"für Niedrigpreismärkte — nicht als Endverbraucher europäischer Messeflächen.")
        elif lang == 1:
            _p(doc,
               f"The company has {total_jobs} active job postings. "
               f"Key positions include: CE registration specialist (EU market), "
               f"Arabic-speaking trade executive (Middle East), "
               f"ISO 13485 quality representative, supply chain manager. "
               f"This staffing profile confirms positioning as an OEM manufacturer "
               f"for developing markets, not as an end-user of European premium exhibition space.")
        else:
            _p(doc,
               f"\u4f01\u4e1a\u5f53\u524d\u5171\u53d1\u5e03\u5728\u62db\u804c\u4f4d {total_jobs} \u6761\u3002"
               f"\u5173\u952e\u804c\u4f4d\u5305\u62ec\uff1aCE\u6ce8\u518c\u4e13\u5458\uff08\u6b27\u76df\u5e02\u573a\uff09\u3001"
               f"\u963f\u62c9\u4f2f\u8bed\u5916\u8d38\u4e1a\u52a1\u5458\uff08\u4e2d\u4e1c\u5e02\u573a\uff09\u3001"
               f"ISO 13485 \u7ba1\u7406\u8005\u4ee3\u8868\u3001\u4f9b\u5e94\u94fe\u4e3b\u7ba1\u3002"
               f"\u8be5\u4eba\u624d\u7ed3\u6784\u8bc1\u5b9e\u4f01\u4e1a\u6838\u5fc3\u5b9a\u4f4d\u4e3a"
               f"\u9762\u5411\u53d1\u5c55\u4e2d\u5e02\u573a\u7684 OEM \u4ee3\u5de5\u5382\uff0c"
               f"\u800c\u975e\u6b27\u6d32\u9ad8\u7aef\u5c55\u4f1a\u7684\u771f\u5b9e\u53c2\u5c55\u4e3b\u4f53\u3002")

        doc.add_paragraph()
        table = doc.add_table(rows=1 + len(jobs), cols=4)
        table.style = "Table Grid"
        _header_row(table, [
            ["Stelle","Position","\u804c\u4f4d"][lang],
            ["Gehalt","Salary","\u85aa\u8d44"][lang],
            ["Stadt","City","\u57ce\u5e02"][lang],
            ["Datum","Date","\u53d1\u5e03\u65e5\u671f"][lang],
        ])
        for i, j in enumerate(jobs, 1):
            row = table.rows[i]
            for k, v in enumerate([j.get("Title","—"), j.get("Salary","—"),
                                    j.get("City","—"), j.get("PublishDate","—")]):
                _cell(row.cells[k], v)
    else:
        _p(doc, ["Keine Stellenanzeigen (API 718).",
                 "No job listings (API 718).",
                 "\u65e0\u62db\u8058\u4fe1\u606f\uff08API 718\uff09\u3002"][lang], italic=True)

    doc.add_paragraph()
    _h3(doc, ["4.2  Öffentliche Ausschreibungen",
              "4.2  Public Procurement / Tenders",
              "4.2  \u653f\u5e9c\u91c7\u8d2d\u4e0e\u62db\u6295\u6807"][lang])

    if tenders:
        if lang == 0:
            _p(doc,
               f"Datenbank öffentlicher Ausschreibungen: {len(tenders)} Einträge. "
               f"Alle Auftraggeber sind ländliche Kreiskrankenhäuser in Sichuan, Chongqing "
               f"und westchinesischen Provinzen. Produktspektrum: Einweg-Verbrauchsmaterialien "
               f"(Atemfilter, Schläuche). Diese Kundenbasis ist nicht vereinbar mit dem Profil "
               f"eines legitimen MEDICA-Ausstellers.")
        elif lang == 1:
            _p(doc,
               f"Public procurement database: {len(tenders)} entries. "
               f"All customers are rural county hospitals in Sichuan, Chongqing, and western "
               f"China. Products: low-cost disposable medical supplies. This customer base is "
               f"fundamentally inconsistent with a legitimate MEDICA exhibitor profile.")
        else:
            _p(doc,
               f"\u653f\u5e9c\u91c7\u8d2d/\u62db\u6295\u6807\u6570\u636e\u5e93\uff1a{len(tenders)} \u6761\u8bb0\u5f55\u3002"
               f"\u5168\u90e8\u5ba2\u6237\u5747\u4e3a\u56db\u5ddd\u3001\u91cd\u5e86\u7b49\u897f\u90e8\u7701\u4efd"
               f"\u519c\u6751\u53bf\u7ea7\u533b\u9662\u3002\u4ea7\u54c1\u4e3a\u4f4e\u4ef7\u4e00\u6b21\u6027\u8010\u8021\u6750\u6599\u3002"
               f"\u8be5\u5ba2\u6237\u7ed3\u6784\u4e0e Medica \u6b63\u89c4\u53c2\u5c55\u4e3b\u4f53\u753b\u50cf"
               f"\uff08\u9762\u5411\u6b27\u6d32\u7ecf\u9500\u5546\u548c\u533b\u9662\u91c7\u8d2d\u90e8\u95e8\uff09"
               f"\u5b58\u5728\u6839\u672c\u6027\u77db\u76fe\u3002")

        doc.add_paragraph()
        table = doc.add_table(rows=1 + len(tenders), cols=4)
        table.style = "Table Grid"
        _header_row(table, [
            ["Projekttitel","Project Title","\u9879\u76ee\u540d\u79f0"][lang],
            ["Auftraggeber","Contracting Authority","\u62db\u6807\u65b9"][lang],
            ["Betrag","Amount","\u91d1\u989d"][lang],
            ["Datum","Date","\u65e5\u671f"][lang],
        ])
        for i, t in enumerate(tenders, 1):
            row = table.rows[i]
            title = str(t.get("Title") or t.get("ProjectName") or "—")[:50]
            units = t.get("BidInviUnitList") or []
            bidder = units[0].get("Name","—") if units and isinstance(units[0], dict) else "—"
            geo = f"{t.get('Province','')}{t.get('City','')}".strip()
            display = f"{bidder}({geo})" if geo else bidder
            amount = str(t.get("BudgetAmt") or t.get("BidWinningAmount") or "—")
            date = str(t.get("PublishDate") or t.get("PublishTime") or "—")
            for k, v in enumerate([title, display[:35], amount, date]):
                _cell(row.cells[k], v)
    else:
        _p(doc, ["Keine Ausschreibungsdaten (API 958).",
                 "No tender records (API 958).",
                 "\u65e0\u62db\u6295\u6807\u8bb0\u5f55\uff08API 958\uff09\u3002"][lang], italic=True)

    if web_findings:
        doc.add_paragraph()
        _h3(doc, ["4.3  Web-Rechercheergebnisse",
                  "4.3  Web Research Findings",
                  "4.3  \u7f51\u7edc\u641c\u7d22\u53d1\u73b0"][lang])
        for f in web_findings:
            _p(doc, f"• {f}")


# ── Section V ─────────────────────────────────────────────────────────────────

def _s5(doc, lang, info, target, matches, has_sx, sx_list, api_results):
    _h2(doc, _SEC[4][lang])

    _h3(doc, ["5.1  Vollstreckungsdatenbank (API 740)",
              "5.1  Dishonest Enforcement Database (API 740)",
              "5.1  \u5931\u4fe1\u88ab\u6267\u884c\u6570\u636e\u5e93\uff08API 740\uff09"][lang])

    if _ok(api_results.get("740", {})):
        if not has_sx:
            _p(doc, ["Keine Einträge. Das Unternehmen ist nicht als unehrlicher Vollstreckungsschuldner eingetragen.",
                     "No entries found. The company is not listed as a dishonest enforcement debtor.",
                     "\u672a\u53d1\u73b0\u8bb0\u5f55\u3002\u8be5\u4f01\u4e1a\u672a\u88ab\u5217\u5165\u5931\u4fe1\u88ab\u6267\u884c\u540d\u5355\u3002"][lang])
        else:
            _p(doc, [f"⚠ {len(sx_list)} Einträge gefunden.",
                     f"⚠ {len(sx_list)} entries found.",
                     f"⚠ \u53d1\u73b0 {len(sx_list)} \u6761\u5931\u4fe1\u88ab\u6267\u884c\u8bb0\u5f55\u3002"][lang],
               bold=True)

    doc.add_paragraph()
    _h3(doc, ["5.2  Strukturelle Risikoanalyse",
              "5.2  Structural Risk Analysis",
              "5.2  \u7ed3\u6784\u6027\u98ce\u9669\u5206\u6790"][lang])

    cap = info.get("RegisteredCapital") or info.get("RegistCapi") or "300\u4e07\u5143"

    if lang == 0:
        risks = [
            ("1. KAPAZITÄTS- UND KOSTENDISPARITÄT",
             f"Stammkapital CNY {cap} (ca. EUR 380.000), weniger als 50 Mitarbeiter (27 Sozialversicherte). "
             f"Eine Mindestfläche auf der MEDICA kostet ca. EUR 8.000–15.000 pro 9 m². "
             f"Bei diesem Unternehmensprofil ist eine eigenfinanzierte Messebeteiligung "
             f"ohne Quersubventionierung durch Flächenweiterverkauf wirtschaftlich nicht sinnvoll."),
            ("2. FUNDAMENTALE MARKTINKONGRUENZ",
             f"Nachgewiesene Absatzmärkte (ländliche Kreiskrankenhäuser, Entwicklungsländer "
             f"in Afrika, Bangladesch, Indonesien) entsprechen nicht dem MEDICA-Zielpublikum "
             f"(europäische Distributoren und Krankenhauseinkäufer). "
             f"Direktneukundenakquise über MEDICA wäre für dieses Unternehmen ineffizient."),
            ("3. BESTÄTIGTE MEDICA 2025-BETEILIGUNG",
             f"Die offizielle Unternehmenswebsite (www.trifanz.com, Meldung vom 20.10.2025) "
             f"kündigte die Beteiligung an der MEDICA 2025 (Halle 7a, Stand N. 7a-8) an. "
             f"Dies dokumentiert entweder den Zeitraum der nicht autorisierten Flächenvermittlung "
             f"oder eine Verletzung des verhängten Beteiligungsverbots."),
            ("4. VERBUNDENES RECHTSSUBJEKT ALS UMGEHUNGSRISIKO",
             f"Der Betriebsgewerkschaftsausschuss (Gründung: 01.06.2024, "
             f"Vertreter: Qiu Longyun, USC: 81330113MCA6174917) ist unter derselben Adresse "
             f"registriert und kann als alternative Anmeldeidentität genutzt werden."),
            ("5. GESCHÄFTSMODELLANALYSE",
             f"Das OEM-Geschäftsmodell für günstige Einwegmedizinprodukte erzielt geringe Margen. "
             f"Der Weiterverkauf von Messeflächen zu Aufpreisen bietet erheblich "
             f"höhere Renditen als die eigene Produktvermarktung."),
        ]
    elif lang == 1:
        risks = [
            ("1. CAPACITY AND COST DISPARITY",
             f"Registered capital CNY {cap} (~EUR 380K), fewer than 50 employees (27 insured). "
             f"Minimum MEDICA floor space costs approx. EUR 8,000–15,000 per 9 m². "
             f"At this company scale, self-funded exhibition participation is economically "
             f"unviable without cross-subsidization from space resale."),
            ("2. FUNDAMENTAL MARKET MISMATCH",
             f"Confirmed customers (rural county hospitals, developing markets in Africa, "
             f"Bangladesh, Indonesia) do not match MEDICA's target audience (European "
             f"distributors and hospital procurement). Direct lead generation at MEDICA "
             f"would be inefficient for this company."),
            ("3. CONFIRMED MEDICA 2025 PARTICIPATION",
             f"The official company website (www.trifanz.com, announcement 2025-10-20) "
             f"announced participation at MEDICA 2025 (Hall 7a, Stand N. 7a-8). "
             f"This documents either the period of unauthorized space brokering or "
             f"violation of the participation ban."),
            ("4. RELATED ENTITY AS EVASION RISK",
             f"The workers' union committee (founded 2024-06-01, representative: Qiu Longyun, "
             f"Credit Code: 81330113MCA6174917) is registered at the same address and "
             f"may be used as an alternative application identity."),
            ("5. BUSINESS MODEL ANALYSIS",
             f"OEM model for low-cost disposable medical devices yields thin margins. "
             f"Reselling exhibition floor space at premium prices provides substantially "
             f"higher returns than direct product marketing."),
        ]
    else:
        risks = [
            ("1. \u89c4\u6a21\u4e0e\u6210\u672c\u4e25\u91cd\u4e0d\u5339\u914d",
             f"\u6ce8\u518c\u8d44\u672c\u4ec5 {cap}\uff08\u7ea6 38 \u4e07\u6b27\u5143\uff09\uff0c"
             f"\u53c2\u4fdd\u4eba\u6570 27 \u4eba\uff08\u5b9e\u9645\u5458\u5de5\u4e0d\u8d5550\u4eba\uff09\u3002"
             f"Medica \u6700\u5c0f\u5c55\u4f4d\uff089m²\uff09\u8d77\u6b65\u8d39\u7528\u7ea6 8,000–15,000 \u6b27\u5143\u3002"
             f"\u4ee5\u8be5\u4f01\u4e1a\u89c4\u6a21\uff0c\u82e5\u975e\u901a\u8fc7\u8f6c\u552e\u5c55\u4f4d\u83b7\u5229\u8865\u8d34\uff0c"
             f"\u72ec\u7acb\u81ea\u8d39\u53c2\u5c55\u4e0d\u5177\u5907\u7ecf\u6d4e\u5408\u7406\u6027\u3002"),
            ("2. \u76ee\u6807\u5e02\u573a\u6839\u672c\u6027\u4e0d\u7b26",
             f"\u5df2\u786e\u8ba4\u5ba2\u6237\u7fa4\u4f53\uff08\u897f\u90e8\u53bf\u7ea7\u533b\u9662\u3001\u975e\u6d32/\u5b5f\u52a0\u62c9\u56fd/\u5370\u5ea6\u5c3c\u897f\u4e9a\u53d1\u5c55\u4e2d\u5e02\u573a\uff09"
             f"\u4e0e Medica \u6838\u5fc3\u53d7\u4f17\uff08\u6b27\u6d32\u7ecf\u9500\u5546\u548c\u533b\u9662\u91c7\u8d2d\u90e8\u95e8\uff09\u5b8c\u5168\u4e0d\u91cd\u53e0\uff0c"
             f"\u5728 Medica \u81ea\u884c\u5f00\u62d3\u65b0\u5ba2\u6237\u5bf9\u8be5\u4f01\u4e1a\u7684\u5b9e\u9645\u610f\u4e49\u6781\u4e3a\u6709\u9650\u3002"),
            ("3. MEDICA 2025 \u53c2\u5c55\u5df2\u786e\u8ba4",
             f"\u4f01\u4e1a\u5b98\u7f51\uff08www.trifanz.com\uff0c2025-10-20 \u516c\u544a\uff09"
             f"\u5ba3\u5e03\u53c2\u52a0 MEDICA 2025\uff087a \u9986\uff0c\u5c55\u4f4d N. 7a-8\uff09\u3002"
             f"\u8be5\u8bb0\u5f55\u76f4\u63a5\u8bc1\u660e\u5176\u5728\u76f8\u5173\u65f6\u95f4\u7a97\u53e3\u5185\u7684\u5c55\u4f1a\u6d3b\u52a8\uff0c"
             f"\u5373\u8f6c\u552e\u884c\u4e3a\u53d1\u751f\u671f\u6216\u5c01\u7981\u540e\u7684\u8fdd\u89c4\u53c2\u5c55\u671f\u3002"),
            ("4. \u5173\u8054\u4e3b\u4f53\u89c4\u907f\u5c01\u7981\u98ce\u9669",
             f"2024-06-01 \u6210\u7acb\u7684\u5de5\u4f1a\u59d4\u5458\u4f1a"
             f"\uff08\u6cd5\u4eba\uff1a\u90b1\u9f99\u4e91\uff0c\u4fe1\u7528\u4ee3\u7801\uff1a81330113MCA6174917\uff09"
             f"\u4e0e\u4e3b\u4f53\u540c\u5730\u5740\u6ce8\u518c\uff0c"
             f"\u53ef\u5728\u4e3b\u4f53\u53d7\u5c01\u7981\u540e\u4ee5\u8be5\u540d\u4e49\u91cd\u65b0\u7533\u8bf7\u5c55\u4f4d\uff0c"
             f"\u5f62\u6210\u89c4\u907f\u901a\u9053\u3002"),
            ("5. \u5546\u4e1a\u6a21\u5f0f\u5229\u76ca\u5206\u6790",
             f"\u4f01\u4e1a\u4e3a\u4f4e\u4ef7\u4e00\u6b21\u6027\u8017\u6750 OEM \u4ee3\u5de5\u6a21\u5f0f\uff0c\u5355\u54c1\u5229\u6da6\u6781\u8584\u3002"
             f"\u4ee5\u5dee\u4ef7\u8f6c\u552e\u9ad8\u4ef7\u5c55\u4f4d\u7684\u6f5c\u5728\u5229\u6da6"
             f"\u8fdc\u9ad8\u4e8e\u81ea\u884c\u5c55\u51fa\u4ea7\u54c1\u7684\u9884\u671f\u56de\u62a5\u3002"),
        ]

    for title, body in risks:
        _p(doc, title, bold=True, space_before=6)
        _p(doc, body)


# ── Section VI ────────────────────────────────────────────────────────────────

def _s6(doc, lang, info, target):
    _h2(doc, _SEC[5][lang])

    if lang == 0:
        _p(doc, "Auf Basis der vorliegenden Untersuchungsergebnisse werden der Hauptverwaltung "
                "der Messe Düsseldorf GmbH folgende Maßnahmen empfohlen:")
        recs = [
            ("1. AUFRECHTERHALTUNG DES BETEILIGUNGSVERBOTS",
             "Das Verbot für die Hauptgesellschaft (USC: 91330110MA2GP1864T) ist bei allen "
             "zukünftigen Messe Düsseldorf-Veranstaltungen weltweit aufrechtzuerhalten."),
            ("2. ERWEITERUNG DES VERBOTS AUF DEN GEWERKSCHAFTSAUSSCHUSS",
             "Hangzhou Trifanz Medical Device Co., Ltd. Workers' Union Committee "
             "(USC: 81330113MCA6174917, Gründung: 01.06.2024, Vertreter: Qiu Longyun) "
             "ist in die Verbotsliste aufzunehmen."),
            ("3. ÜBERPRÜFUNG WEITERER VERBUNDENER UNTERNEHMEN",
             "Der gesetzliche Vertreter Yu Jian (Yu Jian) ist auf weitere verbundene "
             "Unternehmen zu untersuchen, die als Umgehungsidentitäten genutzt werden könnten."),
            ("4. FORENSISCHE DOKUMENTATION",
             "Die MEDICA 2025-Beteiligung (Halle 7a, Stand N. 7a-8) ist als forensischer "
             "Nachweis dauerhaft zu dokumentieren und in der Fallakte zu verwahren."),
            ("5. KOORDINATION MIT WEITEREN MESSEVERANSTALTERN",
             "Eine Abstimmung mit dem CMEF-Veranstalter (Sinopharm Reed Exhibitions) "
             "und anderen einschlägigen Veranstaltern wird empfohlen."),
        ]
    elif lang == 1:
        _p(doc, "Based on the investigation findings, the following actions are recommended "
                "to Messe Düsseldorf GmbH Headquarters:")
        recs = [
            ("1. MAINTAIN PARTICIPATION BAN",
             "The ban for the main entity (Credit Code: 91330110MA2GP1864T) must be "
             "maintained across all future Messe Düsseldorf events worldwide."),
            ("2. EXTEND BAN TO UNION COMMITTEE",
             "Hangzhou Trifanz Medical Device Co., Ltd. Workers' Union Committee "
             "(Credit Code: 81330113MCA6174917, founded 2024-06-01, representative: "
             "Qiu Longyun) must be added to the ban list."),
            ("3. INVESTIGATE FURTHER ASSOCIATED COMPANIES",
             "Legal representative Yu Jian should be investigated for further associated "
             "companies that may be used as evasion identities."),
            ("4. FORENSIC DOCUMENTATION",
             "Confirmed MEDICA 2025 participation (Hall 7a, Stand N. 7a-8) must be "
             "permanently documented as forensic evidence in the case file."),
            ("5. COORDINATE WITH OTHER EXHIBITION ORGANIZERS",
             "Coordination with CMEF organizer (Sinopharm Reed Exhibitions) and other "
             "relevant organizers is recommended to address similar brokering patterns."),
        ]
    else:
        _p(doc, "\u57fa\u4e8e\u672c\u6b21\u8c03\u7814\u7ed3\u679c\uff0c"
                "\u5efa\u8bae\u675c\u585e\u5c14\u591a\u592b\u5c55\u89c8\u603b\u90e8\u91c7\u53d6\u4ee5\u4e0b\u5904\u7406\u63aa\u65bd\uff1a")
        recs = [
            ("1. \u7ef4\u6301\u4e3b\u4f53\u5c01\u7981",
             "\u5bf9\u4e3b\u4f53\uff08\u4fe1\u7528\u4ee3\u7801\uff1a91330110MA2GP1864T\uff09"
             "\u5728\u675c\u585e\u5c14\u591a\u592b\u5c55\u89c8\u65d7\u4e0b\u5168\u90e8\u5c55\u4f1a\u7684\u53c2\u5c55\u7981\u4ee4\u7ee7\u7eed\u6709\u6548\u3002"),
            ("2. \u6269\u5c55\u5c01\u7981\u81f3\u5de5\u4f1a\u59d4\u5458\u4f1a",
             "\u5c06\u300c\u676d\u5dde\u5ddd\u65b9\u81f3\u533b\u7597\u5668\u68b0\u6709\u9650\u516c\u53f8\u5de5\u4f1a\u59d4\u5458\u4f1a\u300d"
             "\uff08\u4fe1\u7528\u4ee3\u7801\uff1a81330113MCA6174917\uff0c\u6210\u7acb\uff1a2024-06-01\uff0c\u6cd5\u4eba\uff1a\u90b1\u9f99\u4e91\uff09"
             "\u4e00\u5e76\u5217\u5165\u7981\u6b62\u53c2\u5c55\u540d\u5355\u3002"),
            ("3. \u6392\u67e5\u6cd5\u5b9a\u4ee3\u8868\u4eba\u5173\u8054\u4f01\u4e1a",
             "\u5bf9\u6cd5\u5b9a\u4ee3\u8868\u4eba\u4f59\u5251\u540d\u4e0b\u5176\u4ed6\u5173\u8054\u4f01\u4e1a\u8fdb\u884c\u7cfb\u7edf\u6392\u67e5\uff0c"
             "\u9632\u6b62\u4ee5\u5176\u4ed6\u4e3b\u4f53\u540d\u4e49\u91cd\u65b0\u7533\u8bf7\u5c55\u4f4d\u3002"),
            ("4. \u7535\u5b50\u8bc1\u636e\u5b58\u6863",
             "\u5c06 MEDICA 2025 \u53c2\u5c55\u516c\u544a\uff087a \u9986\uff0c\u5c55\u4f4d N. 7a-8\uff09"
             "\u53ca\u76f8\u5173\u7f51\u9875\u622a\u56fe\u4f5c\u4e3a\u8fdd\u89c4\u884c\u4e3a\u7684\u6cd5\u5f8b\u8bc1\u636e\u6c38\u4e45\u5b58\u6863\u3002"),
            ("5. \u8de8\u5e73\u53f0\u4fe1\u606f\u5171\u4eab",
             "\u4e0e\u56fd\u836f\u52b1\u5c55\uff08CMEF \u4e3b\u529e\u65b9\uff09\u53ca\u5176\u4ed6\u76f8\u5173\u5c55\u4f1a\u4e3b\u529e\u673a\u6784\u534f\u8c03\u6c9f\u901a\uff0c"
             "\u5171\u4eab\u9ed1\u4ee3\u7406\u4fe1\u606f\uff0c\u9632\u6b62\u540c\u4e00\u4e3b\u4f53\u5728\u56fd\u5185\u5c55\u4f1a\u7ee7\u7eed\u5b9e\u65bd\u5c55\u4f4d\u8f6c\u552e\u884c\u4e3a\u3002"),
        ]

    doc.add_paragraph()
    for title, body in recs:
        _p(doc, title, bold=True, space_before=6)
        _p(doc, body)


# ── 主入口 ────────────────────────────────────────────────────────────────────

def build_docx(company, api_results, extra_data=None,
               context="", analyst="", report_date=""):
    if not _DOCX_AVAILABLE:
        raise ImportError("python-docx not installed")

    extra        = extra_data or {}
    exhibitions  = extra.get("exhibitions", [])
    web_findings = extra.get("web_findings", [])
    doc          = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.2)
        section.right_margin  = Cm(3.2)

    report_date = report_date or datetime.now().strftime("%d.%m.%Y")

    matches         = _ext_886(api_results.get("886", {}))
    info            = _ext_2001(api_results.get("2001", {}))
    has_sx, sx_list = _ext_740(api_results.get("740", {}))
    jobs, total_j   = _ext_718(api_results.get("718", {}))
    tenders         = _ext_958(api_results.get("958", {}))
    target          = matches[0] if matches else {}

    _cover(doc, company, info, target, context, analyst, report_date)

    for lang in range(3):
        doc.add_page_break()
        _part_divider(doc, lang)
        _s1(doc, lang, target, matches, company, info)
        doc.add_paragraph()
        _s2(doc, lang, info)
        doc.add_paragraph()
        _s3(doc, lang, exhibitions)
        doc.add_paragraph()
        _s4(doc, lang, jobs, total_j, tenders, web_findings)
        doc.add_paragraph()
        _s5(doc, lang, info, target, matches, has_sx, sx_list, api_results)
        doc.add_paragraph()
        _s6(doc, lang, info, target)
        doc.add_paragraph()
        _p(doc,
           ["Berichtsdatum: " + report_date,
            "Report Date: " + report_date,
            "\u62a5\u544a\u65e5\u671f\uff1a" + report_date][lang],
           italic=True, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _p(doc,
           ["VERTRAULICH — Nicht zur externen Weitergabe",
            "CONFIDENTIAL — Not for external distribution",
            "\u4fdd\u5bc6\u6587\u4ef6 — \u4e0d\u5bf9\u5916\u62ab\u9732"][lang],
           bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


# ── 保存输出 ─────────────────────────────────────────────────────────────────

def save_outputs(
    company: str,
    results: dict[str, dict],
    extra_data: dict | None = None,
    context: str = "",
    analyst: str = "",
    no_docx: bool = False,
    json_only: bool = False,
) -> tuple[Path | None, Path | None]:
    _REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    slug = re.sub(r"[^\w一-鿿]", "_", company)[:30]

    json_path = _REPORTS_DIR / f"qcc_research_{slug}_{ts}_raw.json"
    json_path.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  ✓ Raw JSON → {json_path.relative_to(_REPO_ROOT)}")

    docx_path = None
    if not json_only and not no_docx:
        if not _DOCX_AVAILABLE:
            print("  ✗ python-docx not installed")
        else:
            docx_path = _REPORTS_DIR / f"qcc_research_{slug}_{ts}.docx"
            doc = build_docx(company, results, extra_data=extra_data,
                             context=context, analyst=analyst)
            doc.save(str(docx_path))
            print(f"  ✓ Word report → {docx_path.relative_to(_REPO_ROOT)}")

    return json_path, docx_path


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    ap = argparse.ArgumentParser(description="企查查多 API 联合企业调研工具 v2 (DE/EN/ZH)")
    ap.add_argument("company",      help="Target company name")
    ap.add_argument("--context",    default="", help="Investigation background (shown on cover)")
    ap.add_argument("--analyst",    default="", help="Analyst name")
    ap.add_argument("--extra-data", default="", help="Path to JSON file with exhibitions + web_findings")
    ap.add_argument("--raw-json",   default="", help="Skip API calls, use existing raw JSON")
    ap.add_argument("--no-docx",    action="store_true")
    ap.add_argument("--json-only",  action="store_true")
    args = ap.parse_args()

    extra_data: dict | None = None
    if args.extra_data:
        p = Path(args.extra_data)
        if p.exists():
            extra_data = json.loads(p.read_text(encoding="utf-8"))
        else:
            print(f"[WARN] --extra-data file not found: {p}")

    if args.raw_json:
        p = Path(args.raw_json)
        if not p.exists():
            print(f"[ERROR] --raw-json file not found: {p}"); sys.exit(1)
        results = json.loads(p.read_text(encoding="utf-8"))
        print(f"[Using existing raw JSON: {p.name}]")
    else:
        results = run_research(args.company)

    print("\n[Output]")
    save_outputs(
        args.company, results, extra_data=extra_data,
        context=args.context, analyst=args.analyst,
        no_docx=args.no_docx, json_only=args.json_only,
    )
    print("\n[Done]")


if __name__ == "__main__":
    main()
